"""
STEP 3 — Secure Text Extraction
Extracts plain text from PDF, DOCX, and TXT files without executing
any embedded code or active content.
"""

import os


def extract(file_path: str) -> tuple[str, str | None]:
    """
    Returns (text, error).
    text  → extracted plain text (may be empty string on failure).
    error → None on success, error message on failure.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)
    else:
        return "", f"No extractor available for extension '{ext}'"


# ── PDF ──────────────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> tuple[str, str | None]:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        parts = []
        for page in doc:
            # get_text("text") returns plain text only — never executes JS/actions
            parts.append(page.get_text("text"))
        doc.close()
        return "\n".join(parts), None
    except Exception as e:
        return "", f"PDF text extraction failed: {e}"


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _extract_docx(file_path: str) -> tuple[str, str | None]:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        paragraphs.append(t)
        return "\n".join(paragraphs), None
    except Exception as e:
        return "", f"DOCX text extraction failed: {e}"


# ── TXT ──────────────────────────────────────────────────────────────────────

def _extract_txt(file_path: str) -> tuple[str, str | None]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                return f.read(), None
        except Exception:
            continue
    return "", "TXT extraction failed: could not decode file with any supported encoding."
